import streamlit as st
import os
import json
from pathlib import Path
from openai import OpenAI
import docx
import PyPDF2
import io
from spire.doc import Document
import tempfile

# --- Configuration & Setup ---
st.set_page_config(page_title="AI 教学助手 - 流行病学与卫生统计", layout="wide")

KNOWLEDGE_DIR = Path(__file__).parent.parent / "knowledge_base"
CASES_DIR = Path(__file__).parent.parent / "cases"
ASSETS_DIR = Path(__file__).parent.parent / "assets"

# --- Configuration & Defaults ---
DEFAULT_API_KEY = os.getenv("DEEPSEEK_API_KEY", "")
DEFAULT_BASE_URL = os.getenv("OPENAI_BASE_URL", "https://api.deepseek.com")
DEFAULT_MODEL_NAME = os.getenv("MODEL_NAME", "deepseek-v4-flash")

def get_api_client():
    # Priority: session_state (if set via secret) > environment variables
    api_key = st.session_state.get("api_key", DEFAULT_API_KEY)
    base_url = st.session_state.get("base_url", DEFAULT_BASE_URL)
    if not api_key:
        st.error("请先在 Zeabur 环境变量中设置 DEEPSEEK_API_KEY。")
        return None
    return OpenAI(api_key=api_key, base_url=base_url)

# --- Sidebar ---
with st.sidebar:
    st.title("📍 导航")
    # Only show these two to students
    menu = st.radio("选择功能", ["知识库问答", "案例模拟训练"])
    
    # Secret access to Teacher Management (append ?admin=true to URL)
    is_admin = st.query_params.get("admin") == "true"
    if is_admin:
        st.divider()
        if st.checkbox("开启管理模式"):
            menu = "教师管理"

# --- Helper Functions ---
def load_knowledge():
    docs = {}
    for file in KNOWLEDGE_DIR.glob("*.md"):
        with open(file, "r", encoding="utf-8") as f:
            docs[file.stem] = f.read()
    return docs

def load_cases():
    cases = []
    for file in CASES_DIR.glob("*.json"):
        try:
            with open(file, "r", encoding="utf-8") as f:
                data = json.load(f)
                data["_filename"] = file.name
                # Ensure essential fields exist
                if "id" not in data:
                    data["id"] = file.stem
                if "title" not in data:
                    data["title"] = f"未命名案例 ({file.stem})"
                cases.append(data)
        except Exception as e:
            # Even if JSON is broken, we list it so it can be deleted
            cases.append({
                "_filename": file.name,
                "id": file.stem,
                "title": f"损坏/格式错误的案例 ({file.name})",
                "error": str(e)
            })
    return cases

def extract_text_from_docx(file):
    doc = docx.Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text_from_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

def extract_text_from_doc(file):
    # spire.doc needs a physical file path or stream
    with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as tmp:
        tmp.write(file.getvalue())
        tmp_path = tmp.name
    
    try:
        doc = Document()
        doc.LoadFromFile(tmp_path)
        text = doc.GetText()
        # Spire.Doc might add an evaluation warning, but for internal teaching it's often acceptable
        # or we can try to strip it if known. 
    finally:
        os.remove(tmp_path)
    return text

# --- Main Logic ---
if menu == "知识库问答":
    st.header("📚 知识库问答")
    st.write("你可以询问传染病临床表现、流行病学特征、诊断方法等内容。")
    
    knowledge_docs = load_knowledge()
    context = "\n\n".join(knowledge_docs.values())
    
    if "qa_messages" not in st.session_state:
        st.session_state.qa_messages = []

    for msg in st.session_state.qa_messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])
            if msg["role"] == "assistant":
                if "[显示猴痘皮疹图]" in msg["content"]:
                    st.image(ASSETS_DIR / "mpox_rash.png", caption="猴痘皮疹典型临床特征参考图")

    if prompt := st.chat_input("输入你的问题..."):
        st.session_state.qa_messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)

        client = get_api_client()
        if client:
            with st.chat_message("assistant"):
                system_prompt = f"""你是一个专业的口岸卫生检疫与现场流行病学教学助手。
请根据以下知识库内容回答学员问题。回答应直接针对问题，不要在结尾添加‘扩展建议’。

关于图片的展示规则：
1. 知识库中包含一张【猴痘皮疹典型临床特征参考图片】。
2. 只有当学员专门询问“猴痘皮疹的形态”、“皮疹特点”或“皮疹演变”等具体细节时，你才可以在回答中包含此标记：[显示猴痘皮疹图]。
3. 严禁在回答“猴痘是什么”或普通症状列表时包含此标记。

知识库内容：
{context}"""
                response = client.chat.completions.create(
                    model=st.session_state.get("model_name", DEFAULT_MODEL_NAME),
                    messages=[
                        {"role": "system", "content": system_prompt},
                        *st.session_state.qa_messages
                    ],
                    stream=True
                )
                full_response = st.write_stream(response)
                
                # Strict trigger: only show if the specific tag is present
                if "[显示猴痘皮疹图]" in full_response:
                    st.image(ASSETS_DIR / "mpox_rash.png", caption="猴痘皮疹典型临床特征参考图")
                
                st.session_state.qa_messages.append({"role": "assistant", "content": full_response})

elif menu == "案例模拟训练":
    st.header("🧪 案例模拟训练")
    cases = load_cases()
    knowledge_docs = load_knowledge()
    context = "\n\n".join(knowledge_docs.values())
    
    if not cases:
        st.warning("暂无案例，请在教师管理中添加。")
    else:
        # Manage case index in session state for linear navigation
        if "current_case_idx" not in st.session_state:
            st.session_state.current_case_idx = 0
        
        # Ensure index is within bounds
        if st.session_state.current_case_idx >= len(cases):
            st.session_state.current_case_idx = 0
            
        case_data = cases[st.session_state.current_case_idx]
        
        st.info(f"**背景：** {case_data.get('background', '暂无背景信息')}")
        
        # Check case format
        if case_data.get("format") == "interactive_v2":
            st.markdown("### 📋 病例信息")
            p_info = case_data["patient_info"]
            st.write(f"**年龄：** {p_info['年龄']}")
            st.write(f"**性别：** {p_info['性别']}")
            st.write(f"**发病时间：** {p_info['发病时间']}")
            st.write(f"**症状：** {', '.join(p_info['症状'])}")
            st.write(f"**旅行史：** {p_info['旅行史']}")
            st.write(f"**接触史：** {p_info['接触史']}")

            st.divider()
            st.markdown("### 🧠 交互式决策判断")
            
            with st.form(f"case_form_{st.session_state.current_case_idx}"):
                user_diseases = st.multiselect("1. 可能的传染病（可多选）", case_data["options"]["possible_diseases"])
                user_measures = st.multiselect("2. 拟采取的处置措施（可多选）", case_data["options"]["measures"])
                user_treatments = st.multiselect("3. 建议的诊断方法（可多选）", case_data["options"]["treatments"])
                
                submit_btn = st.form_submit_button("提交判断并获取 AI 反馈")

            if submit_btn:
                correct = case_data["correct_answers"]
                client = get_api_client()
                if client:
                    with st.spinner("AI 正在分析您的决策..."):
                        user_ans_str = f"""
学员选择：
- 可能疾病：{', '.join(user_diseases)}
- 处置措施：{', '.join(user_measures)}
- 诊断方法：{', '.join(user_treatments)}
"""
                        sys_eval_prompt = f"""你是一个专业的现场流行病学案例导师。
当前案例：{case_data.get('title')}
标准答案：
- 疾病：{', '.join(correct['possible_diseases'])}
- 措施：{', '.join(correct['measures'])}
- 诊断：{', '.join(correct['treatments'])}
参考 SOP：{case_data.get('reference_sop')}

任务：评价学员的选择是否正确、完整，指出错误或缺失，并给出基于 SOP 的专业解析。简洁专业。
"""
                        response = client.chat.completions.create(
                            model=st.session_state.get("model_name", DEFAULT_MODEL_NAME),
                            messages=[
                                {"role": "system", "content": sys_eval_prompt},
                                {"role": "user", "content": f"病例背景：{case_data['background']}\n{user_ans_str}"}
                            ]
                        )
                        feedback = response.choices[0].message.content
                        st.markdown("### 📝 AI 评估反馈")
                        st.write(feedback)
                        st.success(f"**参考依据：** {case_data.get('reference_sop')}")

            # Navigation Button at the bottom
            st.divider()
            col1, col2, col3 = st.columns([2, 2, 2])
            with col2:
                if st.button("进入下一题", use_container_width=True):
                    st.session_state.current_case_idx = (st.session_state.current_case_idx + 1) % len(cases)
                    st.rerun()
            
        else:
            # Original Stages Logic
            if "stages" in case_data:
                current_id = case_data.get("id", "unknown")
                if "sim_case_id" not in st.session_state or st.session_state.sim_case_id != current_id:
                    st.session_state.sim_case_id = current_id
                    st.session_state.sim_step = 0
                    st.session_state.sim_messages = []

                # Current Step Info
                current_step_idx = st.session_state.sim_step
                if current_step_idx < len(case_data["stages"]):
                    stage = case_data["stages"][current_step_idx]
                    st.markdown(f"### 第 {stage['step']} 步：{stage['title']}")
                    st.warning(f"**任务：** {stage['task']}")
                    
                    if stage.get("data"):
                        st.write("**参考数据：**", stage["data"])
                    
                    # Chat for simulation
                    for msg in st.session_state.sim_messages:
                        with st.chat_message(msg["role"]):
                            st.markdown(msg["content"])

                    if sim_prompt := st.chat_input("输入你的回答或想法..."):
                        st.session_state.sim_messages.append({"role": "user", "content": sim_prompt})
                        with st.chat_message("user"):
                            st.markdown(sim_prompt)
                        
                        client = get_api_client()
                        if client:
                            with st.chat_message("assistant"):
                                sys_sim_prompt = f"""你是一个专业的口岸卫生检疫案例教学引导员。当前案例是《{case_data.get('title', '未知')}》。
                        当前阶段：第{stage['step']}步 - {stage['title']}。
                        任务要求：{stage['task']}。
                        引导问题：{', '.join(stage['guiding_questions'])}。

                        请参考以下口岸规范评价学员的回答：
                        {context}

                        评价准则：
                        1. 如果学员回答正确并符合 SOP 要求，请给予肯定并引导其进入下一步。
                        2. 如果回答有疏漏，请指出疏漏点（如：未进行消毒、未核对申明卡等）。
                        3. 不要直接给出答案，要通过追问引导学员思考。
                        4. 保持专业、严谨且具有指导性的语气。
                        """
                                response = client.chat.completions.create(
                                    model=st.session_state.model_name,
                                    messages=[
                                        {"role": "system", "content": sys_sim_prompt},
                                        *st.session_state.sim_messages
                                    ],
                                    stream=True
                                )
                                full_response = st.write_stream(response)
                                st.session_state.sim_messages.append({"role": "assistant", "content": full_response})
                    
                    if st.button("进入下一步") and current_step_idx < len(case_data["stages"]) - 1:
                        st.session_state.sim_step += 1
                        st.session_state.sim_messages = [] 
                        st.rerun()
                else:
                    st.success("🎉 案例模拟已完成！")
                    st.markdown(f"**参考总结：** {case_data.get('answers_summary', '暂无总结')}")
                    if st.button("重置案例"):
                        st.session_state.sim_step = 0
                        st.session_state.sim_messages = []
                        st.rerun()
            else:
                st.error("❌ 该案例数据格式不完整（缺失 stages 字段），无法启动演练。")
                st.info("建议：请前往‘教师管理’移除该案例并重新上传。")

elif menu == "教师管理":
    st.header("👩‍🏫 教师管理端")
    st.write("在这里上传新的知识库文档或案例。")
    
    tab1, tab2 = st.tabs(["知识库文档", "案例库"])
    
    with tab1:
        st.subheader("📁 知识库文档管理")
        docs = load_knowledge()
        if not docs:
            st.info("知识库目前为空。")
        else:
            for doc_name in docs.keys():
                col1, col2 = st.columns([4, 1])
                with col1:
                    st.text(f"📄 {doc_name}.md")
                with col2:
                    if st.button("移除", key=f"del_{doc_name}"):
                        file_path = KNOWLEDGE_DIR / f"{doc_name}.md"
                        if file_path.exists():
                            os.remove(file_path)
                            st.success(f"文档 {doc_name} 已移除")
                            st.rerun()
        
        st.divider()
        st.subheader("📥 上传新文档")
        uploaded_doc = st.file_uploader("支持 Word (.docx, .doc), PDF (.pdf), Markdown (.md)", type=["docx", "doc", "pdf", "md"])
        if uploaded_doc is not None:
            file_extension = uploaded_doc.name.split(".")[-1].lower()
            doc_name = uploaded_doc.name.rsplit(".", 1)[0]
            
            content = ""
            if file_extension == "docx":
                content = extract_text_from_docx(uploaded_doc)
            elif file_extension == "doc":
                content = extract_text_from_doc(uploaded_doc)
            elif file_extension == "pdf":
                content = extract_text_from_pdf(uploaded_doc)
            elif file_extension == "md":
                content = uploaded_doc.read().decode("utf-8")
            
            if st.button(f"保存 {uploaded_doc.name}"):
                with open(KNOWLEDGE_DIR / f"{doc_name}.md", "w", encoding="utf-8") as f:
                    f.write(content)
                st.success(f"文档 {doc_name} 已成功处理并保存为 Markdown。")
                st.rerun()

    with tab2:
        st.subheader("🧪 案例库管理")
        cases = load_cases()
        if not cases:
            st.info("案例库目前为空。")
        else:
            for c in cases:
                cid = c.get('id', 'unknown')
                ctitle = c.get('title', '未命名案例')
                fname = c.get('_filename')
                
                col1, col2 = st.columns([4, 1])
                with col1:
                    st.text(f"📁 {ctitle} (ID: {cid})")
                with col2:
                    # Use filename as a unique key for the button
                    if st.button("移除", key=f"del_f_{fname}"):
                        file_path = CASES_DIR / fname
                        if file_path.exists():
                            os.remove(file_path)
                            st.success(f"案例 {ctitle} 已从服务器移除")
                            st.rerun()
        
        st.divider()
        st.subheader("📤 上传并生成案例")
        st.info("直接上传 SOP、指南或教材，AI 将自动分析文档并生成交互式案例。")
        uploaded_case_doc = st.file_uploader("支持 Word (.docx, .doc), PDF (.pdf), Markdown (.md)", type=["docx", "doc", "pdf", "md"], key="case_uploader")
        
        if uploaded_case_doc is not None:
            file_extension = uploaded_case_doc.name.split(".")[-1].lower()
            doc_name = uploaded_case_doc.name.rsplit(".", 1)[0]
            
            content = ""
            if file_extension == "docx":
                content = extract_text_from_docx(uploaded_case_doc)
            elif file_extension == "doc":
                content = extract_text_from_doc(uploaded_case_doc)
            elif file_extension == "pdf":
                content = extract_text_from_pdf(uploaded_case_doc)
            elif file_extension == "md":
                content = uploaded_case_doc.read().decode("utf-8")

            if st.button(f"生成并保存《{doc_name}》相关案例"):
                client = get_api_client()
                if client:
                    with st.spinner("AI 正在以海关卫生检疫官视角解析文档并设计案例..."):
                        try:
                            gen_prompt = f"""你是一个专业的‘海关卫生检疫业务’教学设计师。
你的身份设定：你是一名在口岸一线工作的海关卫生检疫人员。

请根据以下参考文档，设计一个交互式决策判断案例。

参考文档内容：
{content[:5000]}

设计要求（必须严格遵守）：
1. 身份合规性：海关人员不具备开具处方药、进行临床治疗或进行院内救治的权利。
2. 背景与信息去标签化（极其重要）：严禁在“背景”和“病例信息”字段中出现任何具体的传染病名称（如“猴痘”、“登革热”、“基孔肯雅热”等）。
   - 背景应仅描述：来源地性质（如“某传染病流行区”）、发现方式、客观体征。
   - 目的：确保学员只能通过临床表现和流行病学史进行推断，而不是在背景中直接看到答案。
3. 背景设计准则：背景应仅描述病例发现的客观情境。严禁包含“作为...人员”、“你需要做什么”等任何角色设定。
4. 选项设计禁令：严禁出现‘给予抗生素治疗’、‘开具退烧药’等医疗处方行为。
5. 选项设计准则：处置措施聚焦海关职责（流调、采样、闭环转运等）；诊断方法聚焦实验室检测（RDT、PCR等）。
6. 格式：采用交互式决策判断格式（interactive_v2），返回纯 JSON。

JSON 结构模板：
{{
  "id": "gen_{doc_name[:10]}",
  "format": "interactive_v2",
  "title": "案例标题",
  "background": "简短背景",
  "patient_info": {{
    "年龄": "...",
    "性别": "...",
    "旅行史": "...",
    "症状": ["...", "..."],
    "发病时间": "...",
    "接触史": "..."
  }},
  "options": {{
    "possible_diseases": ["正确项", "干扰项1", "干扰项2", "..."],
    "measures": ["正确项1", "正确项2", "干扰项（海关禁止的操作或错误流程）", "..."],
    "treatments": ["正确诊断方法1", "干扰项（如错误采样或不必要检测）", "..."]
  }},
  "correct_answers": {{
    "possible_diseases": ["正确项"],
    "measures": ["正确项1", "正确项2"],
    "treatments": ["正确诊断方法1"]
  }},
  "reference_sop": "引用文档的具体条款"
}}
"""
                            response = client.chat.completions.create(
                                model=st.session_state.get("model_name", DEFAULT_MODEL_NAME),
                                messages=[{"role": "system", "content": gen_prompt}]
                            )
                            # Strip potential code blocks
                            raw_json = response.choices[0].message.content.strip()
                            if raw_json.startswith("```"):
                                raw_json = raw_json.split("\n", 1)[1].rsplit("\n", 1)[0].strip()
                            
                            case_data = json.loads(raw_json)
                            # Ensure ID is safe
                            cid = case_data.get('id', doc_name)
                            
                            with open(CASES_DIR / f"{cid}.json", "w", encoding="utf-8") as f:
                                json.dump(case_data, f, ensure_ascii=False, indent=2)
                            
                            st.success(f"✅ 案例生成成功！已保存为：{case_data['title']}")
                            st.rerun()
                        except Exception as e:
                            st.error(f"案例生成或解析失败：{e}")
                            st.code(response.choices[0].message.content) # 显示原始回复方便调试
